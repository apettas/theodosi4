from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.auth import login, logout
from django.contrib import messages
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.urls import reverse, reverse_lazy
from django.http import HttpResponse, JsonResponse, FileResponse
from django.db.models import Q, Sum, F, Count
from django.utils import timezone
from django.core.paginator import Paginator
from django.core.exceptions import PermissionDenied, ValidationError
from django.conf import settings

import os
import json
import datetime
import tempfile
from io import BytesIO
import xlsxwriter
import io

# PDF generation imports
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping

# Word generation imports
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Βοηθητική συνάρτηση για τη μετατροπή του request.user σε πραγματικό αντικείμενο User
def get_real_user(request_user):
    """
    Μετατρέπει το request.user (SimpleLazyObject) σε πραγματικό αντικείμενο User
    """
    from django.contrib.auth import get_user_model
    User = get_user_model()
    if request_user.is_authenticated:
        try:
            return User.objects.get(id=request_user.id)
        except User.DoesNotExist:
            # Αν δεν βρεθεί ο χρήστης, επιστρέφουμε None
            return None
    return None
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches

from .models import (
    User, Teacher, Specialty, TeacherSpecialty, Service, 
    SchoolYear, EmployeeType, PYSEEP, ServiceProvider, 
    EmploymentRelation, Application, PriorService, AuditLog
)
from .forms import (
    CustomAuthenticationForm, CustomUserCreationForm, TeacherForm,
    TeacherSpecialtyForm, ApplicationForm, ApplicationStatusForm,
    PriorServiceForm, VerifyPriorServiceForm, ApplicationSearchForm,
    PriorServiceSearchForm, ReportForm, PYSEEPForm,
    PYSEEPServiceReportForm
)

# Βοηθητική συνάρτηση για την καταγραφή ενεργειών
def log_action(request, action, entity, entity_id, old_values=None, new_values=None):
    # Έλεγχος αν ο χρήστης είναι συνδεδεμένος
    if not request.user.is_authenticated:
        return
    
    # Έλεγχος αν ο χρήστης είναι του προσαρμοσμένου μοντέλου User
    from django.contrib.auth import get_user_model
    User = get_user_model()
    
    try:
        # Προσπάθεια εύρεσης του χρήστη στο προσαρμοσμένο μοντέλο User
        custom_user = User.objects.get(username=request.user.username)
        
        AuditLog.objects.create(
            user=custom_user,
            action=action,
            entity=entity,
            entity_id=entity_id,
            ip_address=request.META.get('REMOTE_ADDR'),
            old_values=old_values,
            new_values=new_values,
            session_key=request.session.session_key
        )
    except User.DoesNotExist:
        # Αν ο χρήστης δεν υπάρχει στο προσαρμοσμένο μοντέλο User, δεν καταγράφουμε την ενέργεια
        pass

# Αρχική σελίδα
@login_required
def home(request):
    # Στατιστικά στοιχεία για την αρχική σελίδα
    total_applications = Application.objects.count()
    pending_applications = Application.objects.filter(status__in=['NEW', 'PENDING_DOCS']).count()
    ready_for_pyseep = Application.objects.filter(status='READY_FOR_PYSEEP').count()
    completed_applications = Application.objects.filter(status='COMPLETED').count()
    
    # Πρόσφατες αιτήσεις
    recent_applications = Application.objects.order_by('-created_at')[:10]
    
    # Επερχόμενα ΠΥΣΕΕΠ
    upcoming_pyseep = PYSEEP.objects.filter(date__gte=timezone.now()).order_by('date')[:5]
    
    context = {
        'total_applications': total_applications,
        'pending_applications': pending_applications,
        'ready_for_pyseep': ready_for_pyseep,
        'completed_applications': completed_applications,
        'recent_applications': recent_applications,
        'upcoming_pyseep': upcoming_pyseep,
    }
    
    return render(request, 'proipiresia/home.html', context)

# Σύνδεση χρήστη
def user_login(request):
    if request.user.is_authenticated:
        return redirect('home')
    
    if request.method == 'POST':
        form = CustomAuthenticationForm(data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            log_action(request, 'LOGIN', 'USER', user.id)
            messages.success(request, f'Καλωσήρθατε, {user.get_full_name() or user.username}!')
            return redirect('home')
    else:
        form = CustomAuthenticationForm()
    
    return render(request, 'proipiresia/login.html', {'form': form})

# Αποσύνδεση χρήστη
@login_required
def user_logout(request):
    log_action(request, 'LOGOUT', 'USER', request.user.id)
    logout(request)
    messages.info(request, 'Αποσυνδεθήκατε επιτυχώς.')
    return redirect('login')

# Λίστα εκπαιδευτικών
class TeacherListView(LoginRequiredMixin, ListView):
    model = Teacher
    template_name = 'proipiresia/teacher_list.html'
    context_object_name = 'teachers'
    paginate_by = 20
    
    def get_queryset(self):
        queryset = super().get_queryset()
        search_term = self.request.GET.get('search', '')
        sort_by = self.request.GET.get('sort', 'last_name')
        order = self.request.GET.get('order', 'asc')
        
        if search_term:
            queryset = queryset.filter(
                Q(last_name__icontains=search_term) |
                Q(first_name__icontains=search_term) |
                Q(father_name__icontains=search_term) |
                Q(teacherspecialty__specialty__code__icontains=search_term)
            ).distinct()
        
        # Sorting
        valid_sort_fields = ['last_name', 'first_name', 'father_name', 'phone', 'email']
        if sort_by in valid_sort_fields:
            if order == 'desc':
                sort_by = f'-{sort_by}'
            queryset = queryset.order_by(sort_by)
        else:
            queryset = queryset.order_by('last_name')
        
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['search_term'] = self.request.GET.get('search', '')
        context['current_sort'] = self.request.GET.get('sort', 'last_name')
        context['current_order'] = self.request.GET.get('order', 'asc')
        return context

# Λεπτομέρειες εκπαιδευτικού
class TeacherDetailView(LoginRequiredMixin, DetailView):
    model = Teacher
    template_name = 'proipiresia/teacher_detail.html'
    context_object_name = 'teacher'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['specialties'] = self.object.teacherspecialty_set.all()
        context['applications'] = Application.objects.filter(teacher=self.object).order_by('-created_at')
        
        return context

# Δημιουργία εκπαιδευτικού
class TeacherCreateView(LoginRequiredMixin, CreateView):
    model = Teacher
    form_class = TeacherForm
    template_name = 'proipiresia/teacher_form.html'
    
    def get_success_url(self):
        return reverse('teacher-detail', kwargs={'pk': self.object.pk})
    
    def form_valid(self, form):
        response = super().form_valid(form)
        log_action(
            self.request, 
            'CREATE', 
            'TEACHER', 
            self.object.id,
            new_values=json.dumps({
                'first_name': self.object.first_name,
                'last_name': self.object.last_name,
                'father_name': self.object.father_name,
                'phone': self.object.phone,
                'email': self.object.email
            })
        )
        messages.success(self.request, f'Ο εκπαιδευτικός {self.object} δημιουργήθηκε επιτυχώς.')
        return response

# Ενημέρωση εκπαιδευτικού
class TeacherUpdateView(LoginRequiredMixin, UpdateView):
    model = Teacher
    form_class = TeacherForm
    template_name = 'proipiresia/teacher_form.html'
    
    def get_success_url(self):
        return reverse('teacher-detail', kwargs={'pk': self.object.pk})
    
    def form_valid(self, form):
        old_values = {
            'first_name': self.object.first_name,
            'last_name': self.object.last_name,
            'father_name': self.object.father_name,
            'phone': self.object.phone,
            'email': self.object.email
        }
        
        response = super().form_valid(form)
        
        log_action(
            self.request, 
            'UPDATE', 
            'TEACHER', 
            self.object.id,
            old_values=json.dumps(old_values),
            new_values=json.dumps({
                'first_name': self.object.first_name,
                'last_name': self.object.last_name,
                'father_name': self.object.father_name,
                'phone': self.object.phone,
                'email': self.object.email
            })
        )
        
        messages.success(self.request, f'Ο εκπαιδευτικός {self.object} ενημερώθηκε επιτυχώς.')
        return response

# Προσθήκη ειδικότητας σε εκπαιδευτικό
@login_required
def add_teacher_specialty(request, teacher_id):
    teacher = get_object_or_404(Teacher, pk=teacher_id)
    
    if request.method == 'POST':
        form = TeacherSpecialtyForm(request.POST)
        if form.is_valid():
            specialty = form.save(commit=False)
            specialty.teacher = teacher
            
            # Αν είναι κύρια ειδικότητα, απενεργοποιούμε τις άλλες κύριες
            if specialty.is_primary:
                TeacherSpecialty.objects.filter(teacher=teacher, is_primary=True).update(is_primary=False)
            
            specialty.save()
            
            log_action(
                request, 
                'CREATE', 
                'TEACHER', 
                teacher.id,
                new_values=json.dumps({
                    'specialty': specialty.specialty.code,
                    'is_primary': specialty.is_primary
                })
            )
            
            messages.success(request, f'Η ειδικότητα {specialty.specialty} προστέθηκε επιτυχώς στον εκπαιδευτικό {teacher}.')
            return redirect('teacher-detail', pk=teacher.id)
    else:
        form = TeacherSpecialtyForm()
    
    return render(request, 'proipiresia/teacher_specialty_form.html', {'form': form, 'teacher': teacher})

# Λίστα αιτήσεων
class ApplicationListView(LoginRequiredMixin, ListView):
    model = Application
    template_name = 'proipiresia/application_list.html'
    context_object_name = 'applications'
    paginate_by = 20
    
    def get_queryset(self):
        queryset = super().get_queryset().filter(is_active=True)
        
        # Φιλτράρισμα με βάση τη φόρμα αναζήτησης
        form = ApplicationSearchForm(self.request.GET)
        
        if form.is_valid():
            teacher_name = form.cleaned_data.get('teacher_name')
            school_year = form.cleaned_data.get('school_year')
            current_service = form.cleaned_data.get('current_service')
            status = form.cleaned_data.get('status')
            pyseep = form.cleaned_data.get('pyseep')
            
            if teacher_name:
                queryset = queryset.filter(
                    Q(teacher__last_name__icontains=teacher_name) |
                    Q(teacher__first_name__icontains=teacher_name)
                )
            
            if school_year:
                queryset = queryset.filter(school_year=school_year)
            
            if current_service:
                queryset = queryset.filter(current_service=current_service)
            
            if status:
                queryset = queryset.filter(status=status)
            
            if pyseep:
                queryset = queryset.filter(pyseep=pyseep)
        
        # Sorting
        sort_by = self.request.GET.get('sort', 'created_at')
        order = self.request.GET.get('order', 'desc')
        
        valid_sort_fields = ['id', 'teacher__last_name', 'current_service__name', 'school_year__name', 'employee_type__name', 'status', 'pyseep__act_number', 'created_at']
        if sort_by in valid_sort_fields:
            if order == 'desc':
                sort_by = f'-{sort_by}'
            queryset = queryset.order_by(sort_by)
        else:
            queryset = queryset.order_by('-created_at')
        
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['search_form'] = ApplicationSearchForm(self.request.GET)
        context['current_sort'] = self.request.GET.get('sort', 'created_at')
        context['current_order'] = self.request.GET.get('order', 'desc')
        return context

# Λεπτομέρειες αίτησης
class ApplicationDetailView(LoginRequiredMixin, DetailView):
    model = Application
    template_name = 'proipiresia/application_detail.html'
    context_object_name = 'application'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['prior_services'] = self.object.priorservice_set.all().order_by('start_date', 'end_date')
        
        # Υπολογισμός συνολικής προϋπηρεσίας
        totals = self.object.priorservice_set.aggregate(
            total_years=Sum('years'),
            total_months=Sum('months'),
            total_days=Sum('days')
        )
        
        years = totals['total_years'] or 0
        months = totals['total_months'] or 0
        days = totals['total_days'] or 0
        
        # Κανονικοποίηση
        months += days // 30
        days = days % 30
        years += months // 12
        months = months % 12
        
        context['total_experience'] = {
            'years': years,
            'months': months,
            'days': days
        }
        
        # Έλεγχος για μη ελεγμένες προϋπηρεσίες
        unverified_services = self.object.priorservice_set.filter(
            verified__isnull=True,
            is_active=True
        )
        context['unverified_services_count'] = unverified_services.count()
        context['has_unverified_services'] = unverified_services.exists()
        
        # Φόρμα αλλαγής κατάστασης
        context['status_form'] = ApplicationStatusForm(instance=self.object)
        
        return context

# Δημιουργία αίτησης
class ApplicationCreateView(LoginRequiredMixin, CreateView):
    model = Application
    form_class = ApplicationForm
    template_name = 'proipiresia/application_form.html'
    
    def get_success_url(self):
        return reverse('application-detail', kwargs={'pk': self.object.pk})
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['school_years'] = SchoolYear.objects.all().order_by('-name')
        return context
    
    def form_valid(self, form):
        # Χρήση της βοηθητικής συνάρτησης για τη μετατροπή του request.user
        user = get_real_user(self.request.user)
        form.instance.created_by = user
        form.instance.updated_by = user
        
        response = super().form_valid(form)
        
        log_action(
            self.request, 
            'CREATE', 
            'APPLICATION', 
            self.object.id,
            new_values=json.dumps({
                'teacher': self.object.teacher.id,
                'current_service': self.object.current_service.id,
                'school_year': self.object.school_year.id,
                'employee_type': self.object.employee_type.id,
                'status': self.object.status
            })
        )
        
        messages.success(self.request, f'Η αίτηση για τον εκπαιδευτικό {self.object.teacher} δημιουργήθηκε επιτυχώς.')
        return response

# Ενημέρωση αίτησης
class ApplicationUpdateView(LoginRequiredMixin, UpdateView):
    model = Application
    form_class = ApplicationForm
    template_name = 'proipiresia/application_form.html'
    
    def get_success_url(self):
        return reverse('application-detail', kwargs={'pk': self.object.pk})
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['school_years'] = SchoolYear.objects.all().order_by('-name')
        return context
    
    def form_valid(self, form):
        old_values = {
            'teacher': self.object.teacher.id,
            'current_service': self.object.current_service.id,
            'school_year': self.object.school_year.id,
            'employee_type': self.object.employee_type.id,
            'recruitment_phase': self.object.recruitment_phase,
            'submission_date': self.object.submission_date.strftime('%Y-%m-%d') if self.object.submission_date else None,
            'protocol_number': self.object.protocol_number,
            'pyseep': self.object.pyseep.id if self.object.pyseep else None,
            'status': self.object.status
        }
        
        # Χρήση της βοηθητικής συνάρτησης για τη μετατροπή του request.user
        user = get_real_user(self.request.user)
        form.instance.updated_by = user
        
        response = super().form_valid(form)
        
        new_values = {
            'teacher': self.object.teacher.id,
            'current_service': self.object.current_service.id,
            'school_year': self.object.school_year.id,
            'employee_type': self.object.employee_type.id,
            'recruitment_phase': self.object.recruitment_phase,
            'submission_date': self.object.submission_date.strftime('%Y-%m-%d') if self.object.submission_date else None,
            'protocol_number': self.object.protocol_number,
            'pyseep': self.object.pyseep.id if self.object.pyseep else None,
            'status': self.object.status
        }
        
        log_action(
            self.request, 
            'UPDATE', 
            'APPLICATION', 
            self.object.id,
            old_values=json.dumps(old_values),
            new_values=json.dumps(new_values)
        )
        
        messages.success(self.request, f'Η αίτηση για τον εκπαιδευτικό {self.object.teacher} ενημερώθηκε επιτυχώς.')
        return response

# Διαγραφή αίτησης
class ApplicationDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Application
    template_name = 'proipiresia/application_confirm_delete.html'
    context_object_name = 'application'
    success_url = reverse_lazy('application-list')
    
    def test_func(self):
        """Επιτρέπει διαγραφή όλων των αιτήσεων εκτός από αυτές που είναι 'COMPLETED'"""
        application = self.get_object()
        return application.status != 'COMPLETED'
    
    def handle_no_permission(self):
        """Εμφανίζει μήνυμα σφάλματος αν δεν επιτρέπεται η διαγραφή"""
        messages.error(self.request, 'Οι ολοκληρωμένες αιτήσεις από ΠΥΣΕΕΠ δεν μπορούν να διαγραφούν.')
        return redirect('application-detail', pk=self.get_object().pk)
    
    def delete(self, request, *args, **kwargs):
        """Override για logging της διαγραφής"""
        application = self.get_object()
        
        # Καταγραφή της διαγραφής
        log_action(
            request,
            'DELETE',
            'APPLICATION',
            application.id,
            old_values=json.dumps({
                'teacher': application.teacher.id,
                'current_service': application.current_service.id,
                'school_year': application.school_year.id,
                'status': application.status
            })
        )
        
        messages.success(request, f'Η αίτηση για τον εκπαιδευτικό {application.teacher} διαγράφηκε επιτυχώς.')
        return super().delete(request, *args, **kwargs)

# Αλλαγή κατάστασης αίτησης
@login_required
def change_application_status(request, pk):
    application = get_object_or_404(Application, pk=pk)
    
    if request.method == 'POST':
        form = ApplicationStatusForm(request.POST, instance=application)
        if form.is_valid():
            old_status = application.status
            
            application = form.save(commit=False)
            application.updated_by = get_real_user(request.user)
            application.save()
            
            # Αν η κατάσταση άλλαξε σε "Ολοκληρωμένη από ΠΥΣΕΕΠ", ενημερώνουμε το ιστορικό των προϋπηρεσιών
            if old_status != 'COMPLETED' and application.status == 'COMPLETED' and application.pyseep:
                for service in application.priorservice_set.all():
                    if service.history:
                        service.history += f"\nΕγκρίθηκε από το ΠΥΣΕΕΠ πράξη {application.pyseep.act_number}, ημερομηνία {application.pyseep.date.strftime('%d/%m/%Y')}."
                    else:
                        service.history = f"Εγκρίθηκε από το ΠΥΣΕΕΠ πράξη {application.pyseep.act_number}, ημερομηνία {application.pyseep.date.strftime('%d/%m/%Y')}."
                    
                    service.updated_by = get_real_user(request.user)
                    service.save()
            
            log_action(
                request, 
                'STATUS_CHANGE', 
                'APPLICATION', 
                application.id,
                old_values=json.dumps({'status': old_status}),
                new_values=json.dumps({'status': application.status})
            )
            
            messages.success(request, f'Η κατάσταση της αίτησης άλλαξε σε "{application.get_status_display()}".')
            
            return redirect('application-detail', pk=application.id)
    
    return redirect('application-detail', pk=application.id)

# Δημιουργία νέας έκδοσης αίτησης
@login_required
def create_application_version(request, pk):
    application = get_object_or_404(Application, pk=pk)
    
    if application.status != 'COMPLETED':
        messages.error(request, 'Μόνο ολοκληρωμένες αιτήσεις μπορούν να έχουν νέα έκδοση.')
        return redirect('application-detail', pk=application.id)
    
    new_application = application.create_new_version()
    new_application.updated_by = get_real_user(request.user)
    new_application.save()
    
    log_action(
        request, 
        'CREATE', 
        'APPLICATION', 
        new_application.id,
        old_values=json.dumps({'original_application': application.id}),
        new_values=json.dumps({
            'version': new_application.version,
            'status': new_application.status
        })
    )
    
    messages.success(request, f'Δημιουργήθηκε νέα έκδοση της αίτησης (Έκδοση {new_application.version}).')
    
    return redirect('application-detail', pk=new_application.id)

# Προσθήκη προϋπηρεσίας σε αίτηση
@login_required
def add_prior_service(request, application_id):
    application = get_object_or_404(Application, pk=application_id)
    
    # Έλεγχος αν η αίτηση είναι ολοκληρωμένη
    if application.status == 'COMPLETED':
        messages.error(request, 'Δεν μπορείτε να προσθέσετε προϋπηρεσία σε ολοκληρωμένη αίτηση.')
        return redirect('application-detail', pk=application.id)
    
    if request.method == 'POST':
        form = PriorServiceForm(request.POST)
        if form.is_valid():
            service = form.save(commit=False)
            service.application = application
            user = get_real_user(request.user)
            service.created_by = user
            service.updated_by = user
            
            try:
                # Δεν χρειάζεται να καλέσουμε τη μέθοδο clean() χειροκίνητα
                # καθώς καλείται αυτόματα από τη μέθοδο is_valid() της φόρμας
                service.save()
                
                log_action(
                    request, 
                    'CREATE', 
                    'PRIOR_SERVICE', 
                    service.id,
                    new_values=json.dumps({
                        'application': service.application.id,
                        'service_provider': service.service_provider.id,
                        'start_date': service.start_date.strftime('%Y-%m-%d'),
                        'end_date': service.end_date.strftime('%Y-%m-%d'),
                        'years': service.years,
                        'months': service.months,
                        'days': service.days
                    })
                )
                
                messages.success(request, f'Η προϋπηρεσία προστέθηκε επιτυχώς.')
                return redirect('application-detail', pk=application.id)
            except ValidationError as e:
                for field, errors in e.message_dict.items():
                    for error in errors:
                        form.add_error(field, error)
    else:
        # Φόρτωση προηγούμενων προϋπηρεσιών του εκπαιδευτικού
        previous_services = PriorService.objects.filter(
            application__teacher=application.teacher,
            application__status='COMPLETED',
            is_active=True
        ).exclude(application=application)
        
        if previous_services.exists() and request.GET.get('load_previous') == 'true':
            # Αντιγραφή προηγούμενων προϋπηρεσιών στην τρέχουσα αίτηση
            for prev_service in previous_services:
                # Έλεγχος αν η προϋπηρεσία υπάρχει ήδη στην τρέχουσα αίτηση
                existing = PriorService.objects.filter(
                    application=application,
                    service_provider=prev_service.service_provider,
                    start_date=prev_service.start_date,
                    end_date=prev_service.end_date
                ).exists()
                
                if not existing:
                    new_service = PriorService.objects.create(
                        application=application,
                        service_provider=prev_service.service_provider,
                        protocol_number=prev_service.protocol_number,
                        employment_relation=prev_service.employment_relation,
                        start_date=prev_service.start_date,
                        end_date=prev_service.end_date,
                        years=prev_service.years,
                        months=prev_service.months,
                        days=prev_service.days,
                        reduced_hours=prev_service.reduced_hours,
                        manual_override=prev_service.manual_override,
                        verified=prev_service.verified,
                        history=prev_service.history,
                        notes=prev_service.notes,
                        internal_notes=prev_service.internal_notes,
                        created_by=get_real_user(request.user),
                        updated_by=get_real_user(request.user),
                        version_id=prev_service.version_id
                    )
                    
                    log_action(
                        request, 
                        'CREATE', 
                        'PRIOR_SERVICE', 
                        new_service.id,
                        new_values=json.dumps({
                            'application': new_service.application.id,
                            'service_provider': new_service.service_provider.id,
                            'copied_from': prev_service.id
                        })
                    )
            
            messages.success(request, f'Οι προηγούμενες προϋπηρεσίες φορτώθηκαν επιτυχώς.')
            return redirect('application-detail', pk=application.id)
        
        form = PriorServiceForm()
    
    return render(request, 'proipiresia/prior_service_form.html', {
        'form': form, 
        'application': application,
        'has_previous': PriorService.objects.filter(
            application__teacher=application.teacher,
            application__status='COMPLETED',
            is_active=True
        ).exclude(application=application).exists()
    })

# Επεξεργασία προϋπηρεσίας
@login_required
def edit_prior_service(request, pk):
    service = get_object_or_404(PriorService, pk=pk)
    application = service.application
    
    # Έλεγχος αν η αίτηση είναι ολοκληρωμένη
    if application.status == 'COMPLETED':
        messages.error(request, 'Δεν μπορείτε να επεξεργαστείτε προϋπηρεσία σε ολοκληρωμένη αίτηση.')
        return redirect('application-detail', pk=application.id)
    
    # Έλεγχος αν η προϋπηρεσία έχει επαληθευτεί
    if service.verified and not request.user.is_superuser:
        messages.error(request, 'Δεν μπορείτε να επεξεργαστείτε επαληθευμένη προϋπηρεσία.')
        return redirect('application-detail', pk=application.id)
    
    if request.method == 'POST':
        form = PriorServiceForm(request.POST, instance=service)
        if form.is_valid():
            old_values = {
                'service_provider': service.service_provider.id,
                'protocol_number': service.protocol_number,
                'employment_relation': service.employment_relation.id,
                'start_date': service.start_date.strftime('%Y-%m-%d'),
                'end_date': service.end_date.strftime('%Y-%m-%d'),
                'years': service.years,
                'months': service.months,
                'days': service.days,
                'notes': service.notes,
                'internal_notes': service.internal_notes
            }
            
            service = form.save(commit=False)
            service.updated_by = get_real_user(request.user)
            
            try:
                service.clean()  # Έλεγχος για αλληλεπικαλυπτόμενες περιόδους
                service.save()
                
                new_values = {
                    'service_provider': service.service_provider.id,
                    'protocol_number': service.protocol_number,
                    'employment_relation': service.employment_relation.id,
                    'start_date': service.start_date.strftime('%Y-%m-%d'),
                    'end_date': service.end_date.strftime('%Y-%m-%d'),
                    'years': service.years,
                    'months': service.months,
                    'days': service.days,
                    'notes': service.notes,
                    'internal_notes': service.internal_notes
                }
                
                log_action(
                    request, 
                    'UPDATE', 
                    'PRIOR_SERVICE', 
                    service.id,
                    old_values=json.dumps(old_values),
                    new_values=json.dumps(new_values)
                )
                
                messages.success(request, f'Η προϋπηρεσία ενημερώθηκε επιτυχώς.')
                return redirect('application-detail', pk=application.id)
            except ValidationError as e:
                for field, errors in e.message_dict.items():
                    for error in errors:
                        form.add_error(field, error)
    else:
        form = PriorServiceForm(instance=service)
    
    return render(request, 'proipiresia/prior_service_form.html', {'form': form, 'application': application})

# Διαγραφή προϋπηρεσίας
@login_required
def delete_prior_service(request, pk):
    service = get_object_or_404(PriorService, pk=pk)
    application = service.application
    
    # Έλεγχος αν η αίτηση είναι ολοκληρωμένη
    if application.status == 'COMPLETED':
        messages.error(request, 'Δεν μπορείτε να διαγράψετε προϋπηρεσία από ολοκληρωμένη αίτηση.')
        return redirect('application-detail', pk=application.id)
    
    # Έλεγχος αν η προϋπηρεσία έχει επαληθευτεί
    if service.verified and not request.user.is_superuser:
        messages.error(request, 'Δεν μπορείτε να διαγράψετε επαληθευμένη προϋπηρεσία.')
        return redirect('application-detail', pk=application.id)
    
    if request.method == 'POST':
        service_id = service.id
        service_values = {
            'service_provider': service.service_provider.id,
            'start_date': service.start_date.strftime('%Y-%m-%d'),
            'end_date': service.end_date.strftime('%Y-%m-%d')
        }
        
        service.delete()
        
        log_action(
            request,
            'DELETE',
            'PRIOR_SERVICE',
            service_id,
            old_values=json.dumps(service_values)
        )
        
        messages.success(request, 'Η προϋπηρεσία διαγράφηκε επιτυχώς.')
        return redirect('application-detail', pk=application.id)
    
    return render(request, 'proipiresia/prior_service_confirm_delete.html', {'service': service, 'application': application})

# Επαλήθευση προϋπηρεσίας
@login_required
def verify_prior_service(request, pk):
    service = get_object_or_404(PriorService, pk=pk)
    application = service.application
    
    if request.method == 'POST':
        form = VerifyPriorServiceForm(request.POST)
        if form.is_valid():
            verified = form.cleaned_data.get('verified')
            
            if verified:
                # Επαλήθευση προϋπηρεσίας
                service.mark_as_verified(get_real_user(request.user))
                
                log_action(
                    request,
                    'VERIFY',
                    'PRIOR_SERVICE',
                    service.id,
                    new_values=json.dumps({
                        'verified': service.verified.strftime('%Y-%m-%d %H:%M:%S'),
                        'verified_by': service.verified_by.id
                    })
                )
                
                messages.success(request, 'Η προϋπηρεσία επαληθεύτηκε επιτυχώς.')
            else:
                # Αναίρεση επαλήθευσης
                service.verified = None
                service.verified_by = None
                service.save()
                
                log_action(
                    request,
                    'VERIFY',
                    'PRIOR_SERVICE',
                    service.id,
                    new_values=json.dumps({
                        'verified': None,
                        'verified_by': None
                    })
                )
                
                messages.success(request, 'Η επαλήθευση της προϋπηρεσίας αναιρέθηκε επιτυχώς.')
    
    return redirect('application-detail', pk=application.id)

# Νέα Αναφορά ΠΥΣΕΕΠ - Landscape format
@login_required
def generate_pyseep_service_report(request):
    print(f"DEBUG: generate_pyseep_service_report called with method: {request.method}")
    print(f"DEBUG: GET params: {request.GET}")
    print(f"DEBUG: POST params: {request.POST}")
    
    # Έλεγχος αν έχει περαστεί συγκεκριμένο ΠΥΣΕΕΠ ως παράμετρος (από link)
    pyseep_id = request.GET.get('pyseep')
    if pyseep_id:
        try:
            pyseep = PYSEEP.objects.get(id=pyseep_id)
            # Φιλτράρισμα αιτήσεων για το συγκεκριμένο ΠΥΣΕΕΠ με status 'READY_FOR_PYSEEP' ή 'COMPLETED'
            applications = Application.objects.filter(
                pyseep=pyseep,
                status__in=['READY_FOR_PYSEEP', 'COMPLETED']
            ).select_related(
                'teacher', 'current_service'
            ).prefetch_related(
                'teacher__teacherspecialty_set__specialty',
                'priorservice_set__service_provider',
                'priorservice_set__employment_relation'
            ).order_by('teacher__last_name', 'teacher__first_name')
            
            # Προεπιλεγμένη φόρμα με το ΠΥΣΕΕΠ ήδη επιλεγμένο
            form = PYSEEPServiceReportForm(initial={'pyseep': pyseep})
        except PYSEEP.DoesNotExist:
            form = PYSEEPServiceReportForm()
    elif request.method == 'POST':
        form = PYSEEPServiceReportForm(request.POST)
        if form.is_valid():
            pyseep = form.cleaned_data.get('pyseep')
            report_format = form.cleaned_data.get('format')
            
            # Φιλτράρισμα αιτήσεων για το συγκεκριμένο ΠΥΣΕΕΠ με status 'READY_FOR_PYSEEP' ή 'COMPLETED'
            applications = Application.objects.filter(
                pyseep=pyseep,
                status__in=['READY_FOR_PYSEEP', 'COMPLETED']
            ).select_related(
                'teacher', 'current_service'
            ).prefetch_related(
                'teacher__teacherspecialty_set__specialty',
                'priorservice_set__service_provider',
                'priorservice_set__employment_relation'
            ).order_by('teacher__last_name', 'teacher__first_name')
            
            # Δημιουργία αναφοράς στο επιλεγμένο format (landscape)
            if report_format == 'docx':
                return generate_pyseep_docx_landscape_report(pyseep, applications)
            elif report_format == 'pdf':
                return generate_pyseep_pdf_landscape_report(pyseep, applications)
            else:
                # Fallback σε PDF αν το format δεν αναγνωρίζεται
                return generate_pyseep_pdf_landscape_report(pyseep, applications)
    else:
        form = PYSEEPServiceReportForm()
    
    return render(request, 'proipiresia/pyseep_service_report_form.html', {'form': form})

# Άμεση δημιουργία αναφοράς PDF από link
@login_required
def generate_pyseep_direct_pdf(request, pyseep_id):
    """Άμεση δημιουργία αναφοράς ΠΥΣΕΕΠ σε PDF format από direct link"""
    try:
        pyseep = PYSEEP.objects.get(id=pyseep_id)
        
        # Φιλτράρισμα αιτήσεων για το συγκεκριμένο ΠΥΣΕΕΠ
        applications = Application.objects.filter(
            pyseep=pyseep,
            status__in=['READY_FOR_PYSEEP', 'COMPLETED']
        ).select_related(
            'teacher', 'current_service'
        ).prefetch_related(
            'teacher__teacherspecialty_set__specialty',
            'priorservice_set__service_provider',
            'priorservice_set__employment_relation'
        ).order_by('teacher__last_name', 'teacher__first_name')
        
        # Δημιουργία αναφοράς PDF
        return generate_pyseep_pdf_landscape_report(pyseep, applications)
        
    except PYSEEP.DoesNotExist:
        messages.error(request, f'Το ΠΥΣΕΕΠ με ID {pyseep_id} δεν βρέθηκε.')
        return redirect('home')

# Άμεση δημιουργία αναφοράς Word από link
@login_required
def generate_pyseep_direct_word(request, pyseep_id):
    """Άμεση δημιουργία αναφοράς ΠΥΣΕΕΠ σε Word format από direct link"""
    try:
        pyseep = PYSEEP.objects.get(id=pyseep_id)
        
        # Φιλτράρισμα αιτήσεων για το συγκεκριμένο ΠΥΣΕΕΠ
        applications = Application.objects.filter(
            pyseep=pyseep,
            status__in=['READY_FOR_PYSEEP', 'COMPLETED']
        ).select_related(
            'teacher', 'current_service'
        ).prefetch_related(
            'teacher__teacherspecialty_set__specialty',
            'priorservice_set__service_provider',
            'priorservice_set__employment_relation'
        ).order_by('teacher__last_name', 'teacher__first_name')
        
        # Δημιουργία αναφοράς Word
        return generate_pyseep_docx_landscape_report(pyseep, applications)
        
    except PYSEEP.DoesNotExist:
        messages.error(request, f'Το ΠΥΣΕΕΠ με ID {pyseep_id} δεν βρέθηκε.')
        return redirect('home')

# Δημιουργία landscape αναφοράς Word
def generate_pyseep_docx_landscape_report(pyseep, applications):
    """Δημιουργία landscape αναφοράς ΠΥΣΕΕΠ σε Word format"""
    from docx.shared import Inches, Cm
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    document = Document()
    
    # Ρύθμιση σελίδας σε landscape
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)  # A4 landscape width
    section.page_height = Inches(8.27)  # A4 landscape height
    
    # Μικρότερα περιθώρια για περισσότερο χώρο
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    
    # Ομαδοποίηση εφαρμογών ανά υπηρεσία και φάση πρόσληψης
    grouped_applications = {}
    for application in applications:
        service_name = application.current_service.name if application.current_service else "Άγνωστη Υπηρεσία"
        phase = application.recruitment_phase or "Άγνωστη Φάση"
        key = f"{service_name} - {phase}"
        
        if key not in grouped_applications:
            grouped_applications[key] = []
        grouped_applications[key].append(application)
    
    # Δημιουργία αναφοράς για κάθε ομάδα
    for group_title, group_applications in grouped_applications.items():
        # Τίτλος ομάδας
        title = document.add_heading(group_title, 1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Δημιουργία πίνακα
        # Στήλες: ΗΜ/ΝΙΑ ΥΠΟΒΟΛΗΣ, ΟΝΟΜΑΤΕΠΩΝΥΜΟ, ΚΛΑΔΟΣ, ΕΙΔΟΣ ΠΡΟΥΠΗΡΕΣΙΑΣ, ΑΡ.ΠΡΩΤΟΚΟΛΛΟΥ, ΣΧΕΣΗ ΕΡΓΑΣΙΑΣ, ΑΠΟ, ΕΩΣ, ΕΤΗ, ΜΗΝΕΣ, ΗΜΕΡΕΣ, ΠΑΡΑΤΗΡΗΣΕΙΣ
        table = document.add_table(rows=1, cols=12)
        table.style = 'Table Grid'
        
        # Επικεφαλίδες
        headers = [
            'ΗΜ/ΝΙΑ\nΥΠΟΒΟΛΗΣ\nΑΙΤΗΣΗΣ',
            'ΟΝΟΜΑΤΕΠΩΝΥΜΟ',
            'ΚΛΑΔΟΣ-ΕΙΔΙΚΟΤΗΤΑ',
            'ΕΙΔΟΣ\nΠΡΟΥΠΗΡΕΣΙΑΣ',
            'ΑΡ.\nΠΡΩΤΟΚΟΛΛΟΥ\nΒΕΒ.\nΠΡΟΥΠΗΡΕΣΙΑΣ',
            'ΣΧΕΣΗ ΕΡΓΑΣΙΑΣ',
            'ΧΡ.ΔΙΑΣΤΗΜΑ\nΑΠΟ',
            'ΧΡ.ΔΙΑΣΤΗΜΑ\nΕΩΣ',
            'ΠΡΟΥΠΗΡΕΣΙΑ\nΕΤΗ',
            'ΠΡΟΥΠΗΡΕΣΙΑ\nΜΗΝΕΣ',
            'ΠΡΟΥΠΗΡΕΣΙΑ\nΗΜΕΡΕΣ',
            'ΠΑΡΑΤΗΡΗΣΕΙΣ'
        ]
        
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Προσθήκη ελαφρού γαλάζιου χρώματος στις κεφαλίδες
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shading_elm = parse_xml(r'<w:shd {} w:fill="E0F2F1"/>'.format(nsdecls('w')))
            header_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            
            # Μικρότερη γραμματοσειρά για τις επικεφαλίδες
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Inches(0.1)  # Μικρή γραμματοσειρά
        
        # Ρύθμιση πλάτους στηλών
        widths = [Cm(2), Cm(3), Cm(1.5), Cm(2.5), Cm(2), Cm(2), Cm(1.8), Cm(1.8), Cm(1), Cm(1), Cm(1), Cm(3)]
        for i, width in enumerate(widths):
            if i < len(table.columns):
                table.columns[i].width = width
        
        # Προσθήκη δεδομένων
        for application in group_applications:
            teacher = application.teacher
            
            # Λήψη κύριας ειδικότητας
            specialty_text = "Δεν έχει καταχωρηθεί"  # Default fallback
            
            try:
                primary_specialty = teacher.teacherspecialty_set.filter(is_primary=True).first()
                if not primary_specialty:
                    # Αν δεν υπάρχει κύρια ειδικότητα, πάρε την πρώτη
                    primary_specialty = teacher.teacherspecialty_set.first()
                
                if primary_specialty and primary_specialty.specialty:
                    specialty_text = str(primary_specialty.specialty)
                elif teacher.specialties.exists():
                    # Fallback: πάρε την πρώτη ειδικότητα απευθείας
                    first_specialty = teacher.specialties.first()
                    if first_specialty:
                        specialty_text = str(first_specialty)
            except Exception as e:
                # Αν υπάρχει οποιοδήποτε σφάλμα, χρησιμοποίησε το fallback
                specialty_text = "Σφάλμα ανάκτησης"
            
            # Αν δεν υπάρχουν προϋπηρεσίες, προσθέτουμε μια κενή γραμμή
            prior_services = application.priorservice_set.all().order_by('start_date')
            if not prior_services:
                prior_services = [None]  # Κενή προϋπηρεσία για να εμφανιστεί ο εκπαιδευτικός
            
            # Υπολογισμός συνολικών ετών, μηνών, ημερών
            total_years = application.total_calculated_years
            total_months = application.total_calculated_months
            total_days = application.total_calculated_days
            
            for idx, service in enumerate(prior_services):
                row_cells = table.add_row().cells
                
                # Στην πρώτη γραμμή κάθε εκπαιδευτικού, εμφανίζουμε τα βασικά στοιχεία
                if idx == 0:
                    row_cells[0].text = application.submission_date.strftime('%d/%m/%Y') if application.submission_date else ""
                    row_cells[1].text = f"{teacher.last_name} {teacher.first_name}"
                    row_cells[2].text = specialty_text
                
                # Στοιχεία προϋπηρεσίας (αν υπάρχει)
                if service:
                    row_cells[3].text = service.service_provider.name if service.service_provider else ""
                    row_cells[4].text = service.protocol_number or ""
                    row_cells[5].text = service.employment_relation.name if service.employment_relation else ""
                    row_cells[6].text = service.start_date.strftime('%d/%m/%Y') if service.start_date else ""
                    row_cells[7].text = service.end_date.strftime('%d/%m/%Y') if service.end_date else ""
                    row_cells[8].text = str(service.years) if service.years else "0"
                    row_cells[9].text = str(service.months) if service.months else "0"
                    row_cells[10].text = str(service.days) if service.days else "0"
                    row_cells[11].text = service.notes or ""
                
                # Στυλ κελιών
                for cell in row_cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Inches(0.09)  # Μικρή γραμματοσειρά
            
            # Προσθήκη γραμμής συνόλου για κάθε εκπαιδευτικό
            if prior_services and prior_services[0] is not None:  # Μόνο αν υπάρχουν προϋπηρεσίες
                total_row_cells = table.add_row().cells
                total_row_cells[7].text = "ΣΥΝΟΛΟ:"  # Μετακίνηση στη στήλη ΧΡ.ΔΙΑΣΤΗΜΑ ΕΩΣ
                total_row_cells[8].text = str(total_years)
                total_row_cells[9].text = str(total_months)
                total_row_cells[10].text = str(total_days)
                
                # Προσθήκη ελαφρού γαλάζιου χρώματος στη γραμμή συνόλου
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                
                for i in range(12):  # Όλα τα κελιά της γραμμής
                    cell = total_row_cells[i]
                    # Προσθήκη ελαφρού γαλάζιου χρώματος
                    shading_elm = parse_xml(r'<w:shd {} w:fill="E0F2F1"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # Στυλ για γραμμή συνόλου
                for i in [7, 8, 9, 10]:
                    total_row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in total_row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Inches(0.09)
                
                # Προσθήκη κενής γραμμής μετά το σύνολο κάθε εκπαιδευτικού
                empty_row_cells = table.add_row().cells
                # Κενή γραμμή χωρίς περιεχόμενο για διαχωρισμό
        
        # Προσθήκη κενής γραμμής μετά από κάθε ομάδα
        document.add_paragraph()
            
    
    # Αποθήκευση
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    document.save(temp_file.name)
    
    with open(temp_file.name, 'rb') as file:
        response = HttpResponse(
            file.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="PYSEEP_Report_{pyseep.act_number}_landscape.docx"'
    
    os.unlink(temp_file.name)
    return response

# Δημιουργία landscape αναφοράς PDF
def generate_pyseep_pdf_landscape_report(pyseep, applications):
    """Δημιουργία landscape αναφοράς ΠΥΣΕΕΠ σε PDF format"""
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.units import cm
    from reportlab.platypus import PageBreak
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=1*cm,
        rightMargin=1*cm,
        topMargin=1*cm,
        bottomMargin=1*cm
    )
    
    # Στυλ
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    heading_style = styles['Heading2']
    normal_style = styles['Normal']
    
    story = []
    
    # Ομαδοποίηση εφαρμογών ανά υπηρεσία και φάση πρόσληψης
    grouped_applications = {}
    for application in applications:
        service_name = application.current_service.name if application.current_service else "Άγνωστη Υπηρεσία"
        phase = application.recruitment_phase or "Άγνωστη Φάση"
        key = f"{service_name} - {phase}"
        
        if key not in grouped_applications:
            grouped_applications[key] = []
        grouped_applications[key].append(application)
    
    # Δημιουργία αναφοράς για κάθε ομάδα
    for group_title, group_applications in grouped_applications.items():
        # Τίτλος ομάδας
        story.append(Paragraph(f'<b>{group_title}</b>', heading_style))
        story.append(Paragraph('<br/>', normal_style))
        
        # Δημιουργία δεδομένων πίνακα
        table_data = []
        
        # Επικεφαλίδες
        headers = [
            'ΗΜ/ΝΙΑ\nΥΠΟΒΟΛΗΣ\nΑΙΤΗΣΗΣ',
            'ΟΝΟΜΑΤΕΠΩΝΥΜΟ',
            'ΚΛΑΔΟΣ-ΕΙΔΙΚΟΤΗΤΑ',
            'ΕΙΔΟΣ\nΠΡΟΥΠΗΡΕΣΙΑΣ',
            'ΑΡ.ΠΡΩΤΟΚ.\nΒΕΒ.ΠΡΟΥΠ.',
            'ΣΧΕΣΗ\nΕΡΓΑΣΙΑΣ',
            'ΑΠΟ',
            'ΕΩΣ',
            'ΕΤΗ',
            'ΜΗΝΕΣ',
            'ΗΜΕΡΕΣ',
            'ΠΑΡΑΤΗΡΗΣΕΙΣ'
        ]
        table_data.append(headers)
        
        # Προσθήκη δεδομένων
        for application in group_applications:
            teacher = application.teacher
            
            # Λήψη κύριας ειδικότητας
            specialty_text = "Δεν έχει καταχωρηθεί"  # Default fallback
            
            try:
                primary_specialty = teacher.teacherspecialty_set.filter(is_primary=True).first()
                if not primary_specialty:
                    # Αν δεν υπάρχει κύρια ειδικότητα, πάρε την πρώτη
                    primary_specialty = teacher.teacherspecialty_set.first()
                
                if primary_specialty and primary_specialty.specialty:
                    specialty_text = str(primary_specialty.specialty)
                elif teacher.specialties.exists():
                    # Fallback: πάρε την πρώτη ειδικότητα απευθείας
                    first_specialty = teacher.specialties.first()
                    if first_specialty:
                        specialty_text = str(first_specialty)
            except Exception as e:
                # Αν υπάρχει οποιοδήποτε σφάλμα, χρησιμοποίησε το fallback
                specialty_text = "Σφάλμα ανάκτησης"
            
            # Αν δεν υπάρχουν προϋπηρεσίες, προσθέτουμε μια κενή γραμμή
            prior_services = application.priorservice_set.all().order_by('start_date')
            if not prior_services:
                prior_services = [None]  # Κενή προϋπηρεσία για να εμφανιστεί ο εκπαιδευτικός
            
            # Υπολογισμός συνολικών ετών, μηνών, ημερών
            total_years = application.total_calculated_years
            total_months = application.total_calculated_months
            total_days = application.total_calculated_days
            
            for idx, service in enumerate(prior_services):
                row_data = [''] * 12  # 12 στήλες
                
                # Στην πρώτη γραμμή κάθε εκπαιδευτικού, εμφανίζουμε τα βασικά στοιχεία
                if idx == 0:
                    row_data[0] = application.submission_date.strftime('%d/%m/%Y') if application.submission_date else ""
                    row_data[1] = f"{teacher.last_name} {teacher.first_name}"
                    row_data[2] = specialty_text
                
                # Στοιχεία προϋπηρεσίας (αν υπάρχει)
                if service:
                    row_data[3] = service.service_provider.name if service.service_provider else ""
                    row_data[4] = service.protocol_number or ""
                    row_data[5] = service.employment_relation.name if service.employment_relation else ""
                    row_data[6] = service.start_date.strftime('%d/%m/%Y') if service.start_date else ""
                    row_data[7] = service.end_date.strftime('%d/%m/%Y') if service.end_date else ""
                    row_data[8] = str(service.years) if service.years else "0"
                    row_data[9] = str(service.months) if service.months else "0"
                    row_data[10] = str(service.days) if service.days else "0"
                    row_data[11] = service.notes or ""
                
                table_data.append(row_data)
            
            # Προσθήκη γραμμής συνόλου για κάθε εκπαιδευτικό
            if prior_services and prior_services[0] is not None:  # Μόνο αν υπάρχουν προϋπηρεσίες
                total_row = [''] * 12
                total_row[7] = "ΣΥΝΟΛΟ:"  # Μετακίνηση στη στήλη ΕΩΣ
                total_row[8] = str(total_years)
                total_row[9] = str(total_months)
                total_row[10] = str(total_days)
                table_data.append(total_row)
                
                # Προσθήκη κενής γραμμής μετά το σύνολο κάθε εκπαιδευτικού
                empty_row = [''] * 12
                table_data.append(empty_row)
        
        # Δημιουργία πίνακα
        table = Table(table_data)
        
        # Στυλ πίνακα
        table_style = TableStyle([
            # Στυλ επικεφαλίδων - ελαφρύ γαλάζιο χρώμα (ίδιο με Word)
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.878, 0.949, 0.945)),  # #E0F2F1
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            
            # Στυλ κελιών
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            
            # Στυλ για γραμμές συνόλου
            ('FONTNAME', (7, 1), (7, -1), 'Helvetica-Bold'),  # ΣΥΝΟΛΟ στήλη (μετακινήθηκε στο ΕΩΣ)
            ('FONTNAME', (8, 1), (10, -1), 'Helvetica-Bold'), # Έτη, Μήνες, Ημέρες
        ])
        
        table.setStyle(table_style)
        story.append(table)
        story.append(Paragraph('<br/><br/>', normal_style))
    
    # Κατασκευή PDF
    doc.build(story)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="PYSEEP_Report_{pyseep.act_number}_landscape.pdf"'
    
    return response
        

# AJAX view για δημιουργία νέου φορέα προϋπηρεσίας
@login_required
def create_service_provider_ajax(request):
    if request.method == 'POST':
        name = request.POST.get('name', '').strip()
        
        if not name:
            return JsonResponse({'success': False, 'error': 'Το όνομα του φορέα είναι υποχρεωτικό.'})
        
        # Έλεγχος αν υπάρχει ήδη
        if ServiceProvider.objects.filter(name=name).exists():
            return JsonResponse({'success': False, 'error': 'Ο φορέας υπάρχει ήδη.'})
        
        try:
            # Δημιουργία νέου φορέα
            service_provider = ServiceProvider.objects.create(name=name)
            
            # Καταγραφή της ενέργειας
            log_action(
                request,
                'CREATE',
                'OTHER',
                service_provider.id,
                new_values=json.dumps({'name': name})
            )
            
            return JsonResponse({
                'success': True,
                'id': service_provider.id,
                'name': service_provider.name
            })
            
        except Exception as e:
            return JsonResponse({'success': False, 'error': f'Σφάλμα κατά τη δημιουργία: {str(e)}'})
    
    return JsonResponse({'success': False, 'error': 'Μη έγκυρη μέθοδος.'})


# Δημιουργία ΠΥΣΕΕΠ
class PYSEEPCreateView(LoginRequiredMixin, CreateView):
    model = PYSEEP
    form_class = PYSEEPForm
    template_name = 'proipiresia/pyseep_form.html'
    success_url = reverse_lazy('home')

    def form_valid(self, form):
        response = super().form_valid(form)

        # Καταγραφή της ενέργειας
        log_action(
            self.request,
            'CREATE',
            'OTHER',
            self.object.id,
            new_values=json.dumps({
                'act_number': self.object.act_number,
                'date': self.object.date.strftime('%Y-%m-%d'),
                'school_year': self.object.school_year.name
            })
        )

        messages.success(self.request, f'Το ΠΥΣΕΕΠ {self.object.act_number} δημιουργήθηκε επιτυχώς.')
        return response

# Ενημέρωση ΠΥΣΕΕΠ
class PYSEEPUpdateView(LoginRequiredMixin, UpdateView):
    model = PYSEEP
    form_class = PYSEEPForm
    template_name = 'proipiresia/pyseep_form.html'
    success_url = reverse_lazy('home')

    def form_valid(self, form):
        old_values = {
            'act_number': self.object.act_number,
            'date': self.object.date.strftime('%Y-%m-%d'),
            'school_year': self.object.school_year.name
        }

        response = super().form_valid(form)

        # Καταγραφή της ενέργειας
        log_action(
            self.request,
            'UPDATE',
            'OTHER',
            self.object.id,
            old_values=json.dumps(old_values),
            new_values=json.dumps({
                'act_number': self.object.act_number,
                'date': self.object.date.strftime('%Y-%m-%d'),
                'school_year': self.object.school_year.name
            })
        )

        messages.success(self.request, f'Το ΠΥΣΕΕΠ {self.object.act_number} ενημερώθηκε επιτυχώς.')
        return response


@login_required
def create_pyseep_ajax(request):
    if request.method == 'POST':
        form = PYSEEPForm(request.POST)
        if form.is_valid():
            pyseep = form.save()
            log_action(request, 'CREATE', 'PYSEEP', pyseep.id)
            return JsonResponse({
                'success': True,
                'id': pyseep.id,
                'act_number': pyseep.act_number,
                'date': pyseep.date.strftime('%Y-%m-%d'),
                'school_year': pyseep.school_year.name,
                'display_text': str(pyseep)
            })
        else:
            errors = {}
            for field, messages in form.errors.items():
                errors[field] = '; '.join(messages)
            return JsonResponse({'success': False, 'errors': errors})
    return JsonResponse({'success': False, 'error': 'Μη επιτρεπόμενη μέθοδος.'})




def generate_pyseep_excel_report(pyseep, applications):
    """Δημιουργία αναφοράς PYSEEP σε Excel"""
    import io
    import xlsxwriter
    
    # Δημιουργία του Excel αρχείου
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    worksheet = workbook.add_worksheet('PYSEEP Αναφορά')
    
    # Φόρμα τίτλου
    title_format = workbook.add_format({
        'bold': True,
        'font_size': 16,
        'align': 'center',
        'valign': 'vcenter',
        'fg_color': '#4F81BD',
        'font_color': 'white',
        'border': 1
    })
    
    # Φόρμα επικεφαλίδων ομάδων
    group_header_format = workbook.add_format({
        'bold': True,
        'font_size': 12,
        'align': 'left',
        'valign': 'vcenter',
        'fg_color': '#8EA9DB',
        'border': 1
    })
    
    # Φόρμα επικεφαλίδων πίνακα
    table_header_format = workbook.add_format({
        'bold': True,
        'align': 'center',
        'valign': 'vcenter',
        'fg_color': '#D7E4BC',
        'border': 1
    })
    
    cell_format = workbook.add_format({
        'align': 'left',
        'valign': 'vcenter',
        'border': 1
    })
    
    total_format = workbook.add_format({
        'bold': True,
        'align': 'center',
        'valign': 'vcenter',
        'fg_color': '#FFE699',
        'border': 1
    })
    
    # Τίτλος
    worksheet.merge_range(0, 0, 0, 8, f'ΑΝΑΦΟΡΑ ΠΡΟΥΠΗΡΕΣΙΩΝ - {pyseep.act_number}', title_format)
    
    row = 2
    
    # Επικεφαλίδες πίνακα προϋπηρεσιών
    service_headers = ['Φορέας', 'Αρ. Πρωτοκόλλου', 'Σχέση Εργασίας', 'Από', 'Έως', 'Έτη', 'Μήνες', 'Μέρες', 'Παρατηρήσεις']
    
    # Συλλογή δεδομένων
    applications = Application.objects.filter(
        pyseep=pyseep,
        status__in=['READY_FOR_PYSEEP', 'COMPLETED']
    ).select_related('teacher', 'specialty').prefetch_related('priorservice_set')
    
    # Οργάνωση δεδομένων
    report_data = {}
    
    for app in applications:
        service_name = app.teacher.service.name if app.teacher.service else 'Άγνωστη Υπηρεσία'
        phase_name = app.recruitment_phase.name if app.recruitment_phase else 'Άγνωστη Φάση'
        
        if service_name not in report_data:
            report_data[service_name] = {}
        if phase_name not in report_data[service_name]:
            report_data[service_name][phase_name] = []
        
        # Υπολογισμός στοιχείων εκπαιδευτικού
        teacher_data = {
            'teacher': app.teacher,
            'specialty': app.specialty,
            'submission_date': app.submission_date,
            'submission_comments': app.comments,
            'prior_services': list(app.priorservice_set.all()),
            'total_years': 0,
            'total_months': 0,
            'total_days': 0
        }
        
        # Υπολογισμός συνολικού χρόνου
        for ps in teacher_data['prior_services']:
            teacher_data['total_years'] += ps.years or 0
            teacher_data['total_months'] += ps.months or 0
            teacher_data['total_days'] += ps.days or 0
        
        report_data[service_name][phase_name].append(teacher_data)
    
    # Ταξινόμηση
    for service_name in report_data:
        for phase_name in report_data[service_name]:
            report_data[service_name][phase_name].sort(
                key=lambda x: (x['teacher'].last_name, x['teacher'].first_name)
            )
    
    for service_name, phases in report_data.items():
        # Επικεφαλίδα υπηρεσίας
        worksheet.merge_range(row, 0, row, 8, f'ΥΠΗΡΕΣΙΑ: {service_name}', group_header_format)
        row += 1
        
        for phase_name, teachers in phases.items():
            # Επικεφαλίδα φάσης πρόσληψης
            worksheet.merge_range(row, 0, row, 8, f'ΦΑΣΗ ΠΡΟΣΛΗΨΗΣ: {phase_name}', group_header_format)
            row += 1
            
            for teacher_data in teachers:
                teacher = teacher_data['teacher']
                specialty = teacher_data['specialty']
                
                # Στοιχεία εκπαιδευτικού σε μία γραμμή
                teacher_info = f"{teacher.last_name}, {teacher.first_name}"
                if specialty:
                    teacher_info += f", {specialty.code} - {specialty.description}"
                if teacher_data['submission_date']:
                    teacher_info += f", {teacher_data['submission_date'].strftime('%d/%m/%Y')}"
                if teacher_data['submission_comments']:
                    teacher_info += f", {teacher_data['submission_comments']}"
                
                worksheet.merge_range(row, 0, row, 8, teacher_info, cell_format)
                row += 1
                
                # Επικεφαλίδες πίνακα προϋπηρεσιών
                for col, header in enumerate(service_headers):
                    worksheet.write(row, col, header, table_header_format)
                row += 1
                
                # Προϋπηρεσίες
                for service in teacher_data['prior_services']:
                    worksheet.write(row, 0, service.service_provider or '', cell_format)
                    worksheet.write(row, 1, service.protocol_number or '', cell_format)
                    worksheet.write(row, 2, service.employment_relation or '', cell_format)
                    worksheet.write(row, 3, service.start_date.strftime('%d/%m/%Y') if service.start_date else '', cell_format)
                    worksheet.write(row, 4, service.end_date.strftime('%d/%m/%Y') if service.end_date else '', cell_format)
                    worksheet.write(row, 5, service.years or 0, cell_format)
                    worksheet.write(row, 6, service.months or 0, cell_format)
                    worksheet.write(row, 7, service.days or 0, cell_format)
                    worksheet.write(row, 8, service.notes or '', cell_format)
                    row += 1
                
                # Σύνολο
                worksheet.write(row, 0, 'ΣΥΝΟΛΟ:', total_format)
                worksheet.write(row, 1, '', total_format)
                worksheet.write(row, 2, '', total_format)
                worksheet.write(row, 3, '', total_format)
                worksheet.write(row, 4, '', total_format)
                worksheet.write(row, 5, teacher_data['total_years'], total_format)
                worksheet.write(row, 6, teacher_data['total_months'], total_format)
                worksheet.write(row, 7, teacher_data['total_days'], total_format)
                worksheet.write(row, 8, '', total_format)
                row += 2  # Κενή γραμμή
    
    # Ρυθμίσεις στηλών
    worksheet.set_column(0, 0, 25)  # Φορέας
    worksheet.set_column(1, 1, 15)  # Αρ. Πρωτοκόλλου
    worksheet.set_column(2, 2, 20)  # Σχέση Εργασίας
    worksheet.set_column(3, 4, 12)  # Από, Έως
    worksheet.set_column(5, 7, 8)   # Έτη, Μήνες, Μέρες
    worksheet.set_column(8, 8, 30)  # Παρατηρήσεις
    
    workbook.close()
    output.seek(0)
    
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="pyseep_report_{pyseep.act_number}.xlsx"'
    
    return response



# Missing view functions referenced in urls.py

@login_required
def generate_report(request):
    """Γενικό form για αναφορές"""
    from .forms import ReportForm
    if request.method == 'POST':
        form = ReportForm(request.POST)
        if form.is_valid():
            # Redirect based on report type
            report_type = form.cleaned_data['report_type']
            if report_type == 'pyseep':
                return redirect('generate-pyseep-report')
            elif report_type == 'individual':
                return redirect('individual-service-report')
    else:
        form = ReportForm()
    
    return render(request, 'proipiresia/report_form.html', {
        'form': form,
        'title': 'Επιλογή Τύπου Αναφοράς'
    })



